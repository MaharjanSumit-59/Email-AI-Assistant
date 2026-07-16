import io
import logging

from google.genai import types

logger = logging.getLogger(__name__)

# Mime types Gemini can read natively — sent as inline bytes alongside
# the text prompt.
GEMINI_NATIVE_MIME_TYPES = {
    "application/pdf",
    "image/png",
    "image/jpeg",
    "image/jpg",
    "image/webp",
    "image/heic",
    "image/heif",
}

# .docx has no native Gemini support, so we extract its text ourselves
# and fold it into the prompt instead.
DOCX_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

SUPPORTED_MIME_TYPES = GEMINI_NATIVE_MIME_TYPES | DOCX_MIME_TYPES

# Attachments are read synchronously inside a single AI request, so
# keep both the per-file size and the number of files we bother with
# bounded — otherwise one huge/odd email could blow out the request.
MAX_ATTACHMENT_BYTES = 15 * 1024 * 1024
MAX_ATTACHMENTS_PER_REQUEST = 5
MAX_DOCX_CHARS = 20000


def _extract_docx_text(data: bytes) -> str:
    """
    Flattens a .docx file's paragraphs and tables into plain text.
    """
    try:
        import docx  # python-docx
    except ImportError:
        logger.error(
            "python-docx is not installed; cannot read .docx attachments."
        )
        return ""

    try:
        document = docx.Document(io.BytesIO(data))

        chunks = [p.text for p in document.paragraphs if p.text.strip()]

        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))

        text = "\n".join(chunks).strip()

        if len(text) > MAX_DOCX_CHARS:
            text = text[:MAX_DOCX_CHARS] + "\n...[truncated]"

        return text

    except Exception:
        logger.exception("Failed to parse .docx attachment.")
        return ""


class AttachmentReader:
    """
    Downloads an email's attachments from Gmail and turns the
    readable ones into something Gemini can use:

    - Images / PDFs -> handed to Gemini natively as inline bytes.
    - .docx -> text extracted locally, folded into the text prompt.
    - Anything else (xlsx, zip, unknown types, oversized files, ...)
      -> skipped, and reported back so callers can be honest about
      what wasn't analyzed instead of silently ignoring it.
    """

    def __init__(self, gmail_service):
        self.gmail = gmail_service

    def read(self, message_id, attachments):
        """
        `attachments` is the public attachment list returned by
        GmailService.read_email(): [{attachment_id, filename,
        mime_type, size}, ...].

        Returns:
            parts: list[google.genai.types.Part] — image/PDF content
                to pass straight to Gemini alongside the text prompt.
            attachment_context: str — extracted .docx text to append
                to the text prompt.
            skipped: list[str] — filenames that were not read, either
                because the type isn't supported or the download/parse
                failed.
            analyzed: list[str] — filenames that were successfully
                read, for reporting back to the caller.
        """
        parts = []
        text_chunks = []
        skipped = []
        analyzed = []

        readable = [
            a for a in (attachments or [])
            if a.get("mime_type") in SUPPORTED_MIME_TYPES
        ]

        unsupported = [
            a.get("filename", "attachment")
            for a in (attachments or [])
            if a.get("mime_type") not in SUPPORTED_MIME_TYPES
        ]
        skipped.extend(unsupported)

        for attachment in readable[:MAX_ATTACHMENTS_PER_REQUEST]:
            filename = attachment.get("filename", "attachment")
            mime_type = attachment.get("mime_type", "")
            size = attachment.get("size") or 0

            if size and size > MAX_ATTACHMENT_BYTES:
                skipped.append(filename)
                continue

            # Re-look-up the attachment right before downloading, since
            # Gmail's attachmentId isn't guaranteed stable across calls
            # (see GmailService.find_attachment_meta).
            meta = self.gmail.find_attachment_meta(
                message_id,
                attachment["attachment_id"],
            )

            if not meta:
                skipped.append(filename)
                continue

            try:
                data = self.gmail.get_attachment(
                    message_id,
                    gmail_attachment_id=meta.get("_gmail_attachment_id"),
                    inline_data=meta.get("_inline_data"),
                )
            except Exception:
                logger.exception("Failed to download attachment %s", filename)
                skipped.append(filename)
                continue

            if not data:
                skipped.append(filename)
                continue

            if mime_type in GEMINI_NATIVE_MIME_TYPES:
                parts.append(types.Part.from_bytes(data=data, mime_type=mime_type))
                analyzed.append(filename)
            elif mime_type in DOCX_MIME_TYPES:
                extracted = _extract_docx_text(data)
                if extracted:
                    text_chunks.append(f"--- {filename} ---\n{extracted}")
                    analyzed.append(filename)
                else:
                    skipped.append(filename)

        # Anything past the cap on supported attachments never got a
        # chance to be read at all.
        skipped.extend(
            a.get("filename", "attachment")
            for a in readable[MAX_ATTACHMENTS_PER_REQUEST:]
        )

        attachment_context = "\n\n".join(text_chunks)

        return parts, attachment_context, skipped, analyzed